from pathlib import Path

import pytesseract
from pdf2image import convert_from_path
from docx import Document

from . import utils
from .parser_patient_details import PatientDetailsParser
from .parser_llm import LLMPatientDetailsParser


def extract(
    file_path: str,
    file_format: str,
    use_adaptive_preprocessing: bool = True,
    use_llm_fallback: bool = True,
):
    """
    Extract structured information from PDF or DOCX medical documents.

    Args:
        file_path:
            Path to uploaded document.

        file_format:
            Logical document format:
            "patient_details" or "prescription".

        use_adaptive_preprocessing:
            Use adaptive preprocessing for PDF OCR.
            
        use_llm_fallback:
            Use LLM (Groq) for extraction if regex fails.

    Returns:
        ExtractionResult
    """

    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(
            f"File not found: {file_path}"
        )

    if file_format not in {
        "patient_details",
        "prescription",
    }:
        raise ValueError(
            f"Invalid file format: {file_format}"
        )

    # ========================================================
    # Determine physical document type
    # ========================================================

    extension = file_path.suffix.lower()

    # ========================================================
    # PDF
    # ========================================================

    if extension == ".pdf":

        return _extract_from_pdf(
            file_path,
            file_format,
            use_adaptive_preprocessing,
            use_llm_fallback,
        )

    # ========================================================
    # DOCX
    # ========================================================

    elif extension == ".docx":

        return _extract_from_docx(
            file_path,
            file_format,
            use_llm_fallback,
        )

    else:

        raise ValueError(
            "Unsupported document type. "
            "Only PDF and DOCX are supported."
        )


# ============================================================
# PDF extraction
# ============================================================

def _extract_from_pdf(
    file_path: Path,
    file_format: str,
    use_adaptive_preprocessing: bool,
    use_llm_fallback: bool,
):
    """
    Extract text from PDF using:
        PDF -> image -> preprocessing -> Tesseract OCR
    """

    pages = convert_from_path(
        str(file_path)
    )

    if not pages:
        raise ValueError(
            "No pages found in PDF file"
        )

    document_text = []

    page_count = len(pages)

    # --------------------------------------------------------
    # OCR each page
    # --------------------------------------------------------

    for page_index, page in enumerate(pages):

        try:

            if use_adaptive_preprocessing:

                (
                    scale_factor,
                    block_size,
                    constant,
                ) = utils.get_optimal_preprocessing_params(
                    page
                )

                processed_image = (
                    utils.preprocess_image(
                        page,
                        scale_factor,
                        block_size,
                        constant,
                    )
                )

            else:

                processed_image = (
                    utils.preprocess_image(page)
                )

            text = pytesseract.image_to_string(
                processed_image,
                lang="eng",
                config="--psm 6",
            )

            if text.strip():

                document_text.append(
                    text
                )

        except Exception as exc:

            print(
                f"Warning: Failed to process "
                f"PDF page {page_index + 1}: {exc}"
            )

            continue

    if not document_text:

        raise ValueError(
            "No text extracted from any pages "
            "in the PDF"
        )

    full_text = "\n".join(
        document_text
    )

    return _parse_text(
        full_text,
        file_format,
        document_type="pdf",
        page_count=page_count,
        use_llm_fallback=use_llm_fallback,
    )


# ============================================================
# DOCX extraction
# ============================================================

def _extract_from_docx(
    file_path: Path,
    file_format: str,
    use_llm_fallback: bool,
):
    """
    Extract text directly from a DOCX document.

    DOCX does not require OCR because the document
    normally contains machine-readable text.
    """

    document = Document(
        str(file_path)
    )

    text_parts = []

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:

            text_parts.append(
                text
            )

    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    for table in document.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                cell_text = (
                    cell.text.strip()
                )

                if cell_text:

                    row_text.append(
                        cell_text
                    )

            if row_text:

                text_parts.append(
                    " | ".join(row_text)
                )

    if not text_parts:

        raise ValueError(
            "No text could be extracted "
            "from the DOCX document"
        )

    full_text = "\n".join(
        text_parts
    )

    return _parse_text(
        full_text,
        file_format,
        document_type="docx",
        page_count=None,
        use_llm_fallback=use_llm_fallback,
    )


# ============================================================
# Common parser
# ============================================================

def _parse_text(
    full_text: str,
    file_format: str,
    document_type: str,
    page_count=None,
    use_llm_fallback: bool = True,
):
    """
    Convert extracted text into structured medical data.
    """

    if file_format == "patient_details":

        # Try regex parser first
        regex_result = PatientDetailsParser(
            full_text
        ).parse()
        
        # Check if extraction quality is poor
        high_confidence_count = sum(
            1 for m in regex_result.metadata
            if hasattr(m, 'confidence') and m.confidence == "high"
        )
        
        # If poor quality and LLM fallback enabled, try LLM
        if use_llm_fallback and high_confidence_count < 10:
            try:
                print("Regex extraction poor quality, trying LLM...")
                result = LLMPatientDetailsParser(
                    full_text
                ).parse()
            except Exception as e:
                print(f"LLM extraction failed: {e}, using regex results")
                result = regex_result
        else:
            result = regex_result

    elif file_format == "prescription":

        raise NotImplementedError(
            "Prescription parser not yet implemented"
        )

    else:

        raise ValueError(
            f"Unsupported format: {file_format}"
        )

    # --------------------------------------------------------
    # Attach document information if supported
    # --------------------------------------------------------

    result.document_type = document_type

    # If your ExtractionResult later contains page_count,
    # this can be assigned directly.
    if hasattr(result, "page_count"):

        result.page_count = page_count

    return result