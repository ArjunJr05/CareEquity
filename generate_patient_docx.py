import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_patient_docx(filename, patient_data):
    doc = docx.Document()
    
    # Page setup - Standard 0.75 in margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Styling helpers
    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # 1. Header Banner Title
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_sub = p_header.add_run("CAREEQUITY HEALTH SYSTEM • CLINICAL SUMMARY & DISCHARGE RECORD\n")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(9)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(79, 70, 229) # Primary Indigo

    r_title = p_header.add_run(f"PATIENT DISCHARGE & CLINICAL EVALUATION: {patient_data['name'].upper()}")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42) # Slate Dark

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(14)
    r_div = p_div.add_run("_________________________________________________________________________________")
    r_div.font.color.rgb = RGBColor(226, 232, 240)

    # Helper section header
    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(79, 70, 229)

    # Helper key-value row generator inside table
    def populate_table(table, rows_data):
        for i, (k, v, k2, v2) in enumerate(rows_data):
            row_cells = table.rows[i].cells
            
            # Left key
            row_cells[0].paragraphs[0].text = k
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
            row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(71, 85, 105)
            set_cell_background(row_cells[0], "F8FAFC")
            
            # Left val
            row_cells[1].paragraphs[0].text = str(v)
            row_cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
            row_cells[1].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(15, 23, 42)
            
            # Right key
            if k2:
                row_cells[2].paragraphs[0].text = k2
                row_cells[2].paragraphs[0].runs[0].font.bold = True
                row_cells[2].paragraphs[0].runs[0].font.size = Pt(9.5)
                row_cells[2].paragraphs[0].runs[0].font.name = 'Arial'
                row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(71, 85, 105)
                set_cell_background(row_cells[2], "F8FAFC")
                
                # Right val
                row_cells[3].paragraphs[0].text = str(v2)
                row_cells[3].paragraphs[0].runs[0].font.size = Pt(9.5)
                row_cells[3].paragraphs[0].runs[0].font.name = 'Arial'
                row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(15, 23, 42)

            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, 80, 80, 120, 120)

    # 2. PATIENT DEMOGRAPHICS SECTION
    add_section_header("1. PATIENT DEMOGRAPHICS & LIFESTYLE")
    t_demo = doc.add_table(rows=5, cols=4)
    t_demo.alignment = WD_TABLE_ALIGNMENT.CENTER
    demo_data = [
        ("Patient Name |", f"{patient_data['name']} |", "UHID / MRN |", patient_data["mrn"]),
        ("Age / Sex |", f"{patient_data['age']} Years / {patient_data['gender']}", "DOB |", patient_data["dob"]),
        ("Race & Ethnicity |", patient_data.get("race_ethnicity", "Asian / Pacific Islander"), "Smoking Status |", patient_data.get("smoking_status", "Never Smoker")),
        ("Alcohol Use |", patient_data.get("alcohol_use", "None"), "Daily Sedentary Time |", f"{patient_data.get('sedentary_time_min', 360)} min/day"),
        ("Target Locations |", str(patient_data.get("target_locations", [["Trego County", "Kansas", "United States"]])), "Phone |", patient_data["phone"])
    ]
    populate_table(t_demo, demo_data)

    # 3. VITAL SIGNS & PHYSICAL MEASUREMENTS
    add_section_header("2. VITAL SIGNS & ANTHROPOMETRICS")
    t_vitals = doc.add_table(rows=4, cols=4)
    t_vitals.alignment = WD_TABLE_ALIGNMENT.CENTER
    vitals_data = [
        ("Height |", f"{patient_data['height_cm']} cm", "Weight |", f"{patient_data['weight_kg']} kg"),
        ("BMI |", f"{patient_data['bmi']}", "Waist Circumference |", f"{patient_data.get('waist_cm', 88)} cm"),
        ("BP |", f"{patient_data['blood_pressure']}", "Pulse / Heart Rate |", f"{patient_data['pulse']} bpm"),
        ("Temp |", f"{patient_data['temperature']}°F", "SpO2 |", f"{patient_data['spo2']}%")
    ]
    populate_table(t_vitals, vitals_data)

    # 4. METABOLIC & LIVER LAB PANELS
    add_section_header("3. METABOLIC & LIVER LABORATORY PANELS")
    t_labs = doc.add_table(rows=5, cols=4)
    t_labs.alignment = WD_TABLE_ALIGNMENT.CENTER
    labs_data = [
        ("HbA1c |", f"{patient_data.get('hba1c', 6.5)}%", "Fasting Glucose |", f"{patient_data.get('fasting_glucose', 110)} mg/dL"),
        ("Total Cholesterol |", f"{patient_data.get('total_cholesterol', 195)} mg/dL", "LDL Cholesterol |", f"{patient_data.get('ldl', 115)} mg/dL"),
        ("HDL Cholesterol |", f"{patient_data.get('hdl', 48)} mg/dL", "Triglycerides |", f"{patient_data.get('triglycerides', 150)} mg/dL"),
        ("ALT |", f"{patient_data.get('alt', 24)} U/L", "AST |", f"{patient_data.get('ast', 22)} U/L"),
        ("Albumin |", f"{patient_data.get('albumin', 4.2)} g/dL", "Total Bilirubin |", f"{patient_data.get('bilirubin', 0.8)} mg/dL")
    ]
    populate_table(t_labs, labs_data)

    # 5. CHRONIC MEDICAL CONDITIONS
    add_section_header("4. CLINICAL CONTEXT & CHRONIC CONDITIONS")
    t_cond = doc.add_table(rows=2, cols=4)
    t_cond.alignment = WD_TABLE_ALIGNMENT.CENTER
    cond_data = [
        ("Diabetes |", patient_data["diabetes"], "Hypertension |", patient_data["hypertension"]),
        ("Heart Disease |", patient_data["heart_disease"], "Asthma |", patient_data["asthma"])
    ]
    populate_table(t_cond, cond_data)

    # 6. CLINICAL NOTES & ADVICE
    add_section_header("5. CHIEF COMPLAINT & DISCHARGE ADVICE")
    p_notes = doc.add_paragraph()
    p_notes.paragraph_format.space_before = Pt(4)
    r_cc_hdr = p_notes.add_run("Chief Complaint | ")
    r_cc_hdr.font.bold = True
    r_cc_hdr.font.size = Pt(10)
    r_cc_hdr.font.color.rgb = RGBColor(79, 70, 229)
    r_cc_val = p_notes.add_run(patient_data["chief_complaint"] + "\n\n")
    r_cc_val.font.size = Pt(10)

    r_adv_hdr = p_notes.add_run("Discharge Advice & Clinical Summary:\n")
    r_adv_hdr.font.bold = True
    r_adv_hdr.font.size = Pt(10)
    r_adv_hdr.font.color.rgb = RGBColor(15, 23, 42)

    r_adv_val = p_notes.add_run(patient_data["notes"])
    r_adv_val.font.size = Pt(9.5)
    r_adv_val.font.color.rgb = RGBColor(51, 65, 85)

    doc.save(filename)
    print(f"[OK] Generated complete 23-feature clinical test DOCX: {filename}")

if __name__ == "__main__":
    patient_1 = {
        "name": "Veena Kumari",
        "mrn": "UHID-9821435",
        "age": 54,
        "gender": "Female",
        "dob": "14/08/1972",
        "phone": "9840123456",
        "email": "veena.kumari@example.com",
        "address": "No. 42 Anna Nagar, Chennai",
        "race_ethnicity": "Asian / South Asian",
        "smoking_status": "Never Smoker",
        "alcohol_use": "None",
        "sedentary_time_min": 320,
        "target_locations": [["Limestone County", "Alabama", "United States"], ["Adams County", "Ohio", "United States"]],
        "height_cm": 162,
        "weight_kg": 68,
        "bmi": 25.9,
        "waist_cm": 84,
        "blood_pressure": "140/90",
        "pulse": 82,
        "temperature": 98.6,
        "spo2": 98,
        "rr": 18,
        "hba1c": 8.2,
        "fasting_glucose": 142,
        "total_cholesterol": 215,
        "ldl": 130,
        "hdl": 45,
        "triglycerides": 180,
        "alt": 28,
        "ast": 26,
        "albumin": 4.1,
        "bilirubin": 0.7,
        "diabetes": "Yes (HbA1c 8.2%, Positive)",
        "hypertension": "Yes (BP 140/90 mmHg, High)",
        "heart_disease": "No",
        "asthma": "No",
        "chief_complaint": "Type 2 Diabetes Mellitus with Essential Hypertension evaluation",
        "notes": "Patient presented with elevated fasting blood glucose levels and mild fatigue. Recommended diabetic dietary management, daily blood pressure tracking, and regular moderate exercise. Follow up in 4 weeks."
    }
    
    patient_2 = {
        "name": "Robert Chen",
        "mrn": "MRN-5542109",
        "age": 62,
        "gender": "Male",
        "dob": "03/11/1963",
        "phone": "9175550192",
        "email": "robert.chen@example.com",
        "address": "742 Evergreen Terrace, Boston, MA",
        "race_ethnicity": "Asian / East Asian",
        "smoking_status": "Former Smoker",
        "alcohol_use": "Moderate (1-2 drinks/week)",
        "sedentary_time_min": 450,
        "target_locations": [["Trego County", "Kansas", "United States"]],
        "height_cm": 178,
        "weight_kg": 85,
        "bmi": 26.8,
        "waist_cm": 96,
        "blood_pressure": "150/95",
        "pulse": 88,
        "temperature": 98.4,
        "spo2": 95,
        "rr": 20,
        "hba1c": 6.1,
        "fasting_glucose": 108,
        "total_cholesterol": 240,
        "ldl": 155,
        "hdl": 42,
        "triglycerides": 210,
        "alt": 35,
        "ast": 31,
        "albumin": 4.3,
        "bilirubin": 0.9,
        "diabetes": "No",
        "hypertension": "Yes (150/95 mmHg, High)",
        "heart_disease": "Yes (Coronary Artery Disease Present)",
        "asthma": "Yes (Asthmatic symptoms present, 45%)",
        "chief_complaint": "Shortness of breath and chest tightness on exertion",
        "notes": "Patient has a history of bronchial asthma and hypertension. EKG shows sinus rhythm with minor ST alterations. Advised continuation of daily bronchodilator inhaler and low-sodium diet."
    }

    create_patient_docx("Discharge_Summary_Veena_Kumari_Apollo_Diabetes.docx", patient_1)
    create_patient_docx("Discharge_Summary_Robert_Chen_Clinical_Record.docx", patient_2)
